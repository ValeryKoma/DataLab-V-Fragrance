# -*- coding: utf-8 -*-
"""
Bouwt technisch_rapport.docx: het bestaande rapport + alle toevoegingen
(Implementatie 3, fine-tuning, hybride retrieval, webapp, resultaten, discussie).
Bewust ONGESTYLED: alleen vette koppen + echte tabellen. Styling doet de gebruiker zelf.
"""
from docx import Document
from docx.shared import Pt

doc = Document()


def title(t):
    p = doc.add_paragraph()
    r = p.add_run(t)
    r.bold = True
    r.font.size = Pt(14)


def h1(t):
    p = doc.add_paragraph()
    p.add_run(t).bold = True


def h2(t):
    p = doc.add_paragraph()
    p.add_run(t).bold = True


def para(t):
    doc.add_paragraph(t)


def bullet(t):
    doc.add_paragraph(t, style="List Bullet")


def mono(t):
    p = doc.add_paragraph()
    r = p.add_run(t)
    r.font.name = "Consolas"
    r.font.size = Pt(9)


def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        run = c.paragraphs[0].add_run(h)
        run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    doc.add_paragraph("")


# ============================================================ TITEL
title("An AI Recommendation Engine for Personalized Fragrance Discovery for Fragheadsunited")
para("M. J. Reszka-Gniecki, J. Debi-Tewari, V. S. Komarov")
para("Datalab V - Haagse Hogeschool, Zoetermeer, The Netherlands")

# ============================================================ ABSTRACT
h1("Abstract")
para("De parfumindustrie groeit richting personalisatie, maar bestaande platforms zoals "
     "Fragrantica en Parfumo bieden geen intelligente afstemming op individuele smaak - zij "
     "vertrouwen op community beoordelingen en handmatige filters. Dit project, uitgevoerd in "
     "opdracht van startup Fragheadsunited, onderzoekt hoe een Retrieval-Augmented Generation "
     "(RAG) pipeline gepersonaliseerde geuraanbevelingen kan genereren op basis van natural "
     "language input. Het onderzoek is iteratief verlopen: twee verkennende implementaties "
     "(ChromaDB met all-MiniLM-L6-v2; vectordb2 met BGE-small) zijn gebouwd en geevalueerd om de "
     "haalbaarheid aan te tonen, waarna op basis van hun beperkingen een uiteindelijk, "
     "aanmerkelijk beter presterend systeem is ontwikkeld. Dat eindsysteem - de Fragrance "
     "Advisor - gebruikt een op het geurdomein fine-getuned BGE-base embeddingmodel, hybride "
     "retrieval (noten + reviews + lexicaal), een persoonlijk smaakprofiel en een conversationele "
     "LLM-laag. Er is een dataset van 84.144 parfums samengesteld uit drie bronnen, verrijkt met "
     "Parfumo-ratings en gefilterd op datakwaliteit. Het systeem is geevalueerd met Precision@5, "
     "een held-out Recall@10 en kwalitatieve inspectie; het behaalt een gemiddelde Precision@5 "
     "van 0,92 op een set thematische queries en genereert relevante aanbevelingen die loskomen "
     "van populaire communityratings, wat de meerwaarde van een vectorgebaseerde, "
     "gepersonaliseerde aanpak ten opzichte van traditionele filtering bevestigt.")

# ============================================================ I. INLEIDING
h1("I. Inleiding")
para("De parfumindustrie is een omvangrijke en groeiende markt. Alleen al in Nederland genereert "
     "de geursector jaarlijks meer dan EUR 400 miljoen aan omzet [1]. Tegelijkertijd verschuift "
     "het consumentengedrag richting bewustere en persoonlijkere aankooppatronen, waarbij "
     "consumenten geuren zoeken die aansluiten bij hun emoties en persoonlijkheid in plaats van "
     "trends te volgen [2]. Het premiumsegment van de geurenmarkt groeide in 2023 met 12% op "
     "jaarbasis, wat deze verschuiving naar niche en personalisatie onderstreept [3].")
para("Ondanks deze groeiende vraag ontbreekt het aan een intelligent systeem dat individuele "
     "smaakvoorkeuren en emotionele associaties kan vertalen naar een gepersonaliseerde "
     "geuraanbeveling. Huidige platforms zoals Fragrantica en Parfumo bieden uitgebreide "
     "databases met meer dan 100.000 geuren, maar zijn afhankelijk van communitybeoordelingen en "
     "handmatige zoekfilters. Dit leidt tot wat in dit project het rating-paradox wordt genoemd: "
     "een geur met een beoordeling van 9/10 door 10.000 gebruikers verschijnt prominent, terwijl "
     "een geur met een 4/10 effectief onzichtbaar is - ook als deze de beste match zou zijn voor "
     "een specifieke gebruiker.")
para("De opkomst van Retrieval-Augmented Generation (RAG) biedt nieuwe mogelijkheden om dit "
     "probleem aan te pakken. RAG combineert de generatieve capaciteiten van grote taalmodellen "
     "(LLMs) met een extern geheugen in de vorm van een vectordatabase, waardoor feitelijk "
     "nauwkeurigere en contextbewuste antwoorden gegenereerd kunnen worden [4]. Tegelijkertijd "
     "maken sentence-embedding modellen zoals all-MiniLM-L6-v2 het mogelijk om natuurlijke taal "
     "semantisch te representeren in een vectorruimte [5]. Deze technologieen zijn pas sinds "
     "2023-2024 op productieniveau beschikbaar onder open-source licenties (MIT), wat dit project "
     "nu pas technisch haalbaar maakt.")
para("Dit project is uitgevoerd in opdracht van Fragheadsunited, een startup die een "
     "aanbevelingsplatform voor geuren ontwikkelt. Het doel is het bouwen van een "
     "proof-of-concept chatbot die op basis van natural language input gepersonaliseerde "
     "geuraanbevelingen genereert via een RAG-pipeline. Hierbij beschrijft de gebruiker zijn "
     "geurvoorkeuren in natuurlijke taal (bijv. \"iets warms en houtigs voor een date night\"), "
     "waarna het systeem de intentie extraheert, zoekt in een vectordatabase van geindexeerde "
     "geuren, en een gepersonaliseerde aanbeveling retourneert - ongeacht de communautaire "
     "beoordeling van die geur.")
para("De hoofdonderzoeksvraag is: Hoe kan een RAG-gebaseerd AI-aanbevelingssysteem "
     "gepersonaliseerde geuraanbevelingen genereren op basis van natural language "
     "gebruikersvoorkeuren en productintelligentie?")
para("Om deze vraag te beantwoorden zijn drie deelvragen geformuleerd die de structuur van het "
     "onderzoek bepalen: (i) Hoe kan uit verspreide, onvolledige webbronnen een kwalitatief "
     "bruikbare geurdataset worden samengesteld? (ii) Welke combinatie van embeddingmodel en "
     "vectordatabase levert de meest relevante aanbevelingen voor vrije-tekstqueries? (iii) In "
     "hoeverre wijken de aanbevelingen van het systeem af van een aanpak die puur op "
     "communityratings sorteert?")
para("Dit paper is als volgt opgebouwd: Sectie II bespreekt gerelateerd werk. Sectie III "
     "beschrijft de gebruikte data. Sectie IV presenteert de methode en systeemarchitectuur, "
     "inclusief de implementaties die we hebben onderzocht. Sectie V toont de resultaten, en "
     "Sectie VI bevat de discussie en conclusie.")

# ============================================================ II. GERELATEERD WERK
h1("II. Gerelateerd Werk")
h2("A. Aanbevelingssystemen en het Geurdomein")
para("Aanbevelingssystemen worden traditioneel ingedeeld in content-based filtering, "
     "collaborative filtering en hybride benaderingen [6]. Content-based filtering matcht "
     "item-eigenschappen met gebruikersprofielen, terwijl collaborative filtering patronen in "
     "gebruikersgedrag benut. Binnen het geurdomein zijn aanbevelingssystemen relatief "
     "onderontwikkeld. Demarqui [7] ontwikkelde een perfume recommender dat Fragrantica-data "
     "scrapt en cosine similarity op geurnoot-vectoren berekent. Hoewel functioneel, ontbreekt "
     "semantisch begrip van gebruikersbeschrijvingen en kan het geen natuurlijke taalinput "
     "verwerken. Gonzales Martinez [8] ontwikkelde een Bayesiaans hierarchisch model dat "
     "geurvoorkeuren koppelt aan persoonlijkheidsarchetypen, maar implementeerde dit niet als "
     "interactief prototype.")
h2("B. Retrieval-Augmented Generation (RAG)")
para("Het RAG-paradigma werd geintroduceerd door Lewis et al. [4] en combineert een "
     "voorgetraind taalmodel met een dense vector index, waarbij een neurale retriever relevante "
     "documenten ophaalt die de generator vervolgens gebruikt om antwoorden te produceren. "
     "RAG-modellen presteren aantoonbaar beter dan puur parametrische modellen op "
     "kennisintensieve taken [4]. Recente toepassingen tonen de brede inzetbaarheid aan: Yousefi "
     "Maragheh et al. [9] introduceerden een agentic RAG-framework voor gepersonaliseerde "
     "aanbevelingen dat tot 42,1% verbetering in NDCG@5 behaalde ten opzichte van standaard "
     "RAG-baselines. Tot op heden is RAG echter niet toegepast op het geurdomein.")
h2("C. Sentence Embeddings en Olfactorische Taal")
para("Het vertalen van subjectieve geurbeschrijvingen naar machine-leesbare representaties "
     "vereist modellen die de semantiek van natuurlijke taal vastleggen. Sentence-BERT (SBERT) "
     "maakt gebruik van een Siamese BERT-architectuur om zinnen te encoderen in een dichte "
     "vectorruimte [5]. Het all-MiniLM-L6-v2 model genereert 384-dimensionale embeddings en is "
     "bijzonder geschikt voor semantische zoektaken. Nieuwere embeddingfamilies zoals BGE [11] "
     "scoren hoger op de Massive Text Embedding Benchmark (MTEB) [12], wat ons aanleiding gaf "
     "beide modellen te vergelijken (Sectie IV).")
para("Een kernuitdaging is dat geur de minst gelexicaliseerde zintuiglijke modaliteit is: "
     "slechts circa 1% van Engelse woorden heeft olfactorische associaties [10]. Kurfali et al. "
     "[10] onderzochten systematisch of taalmodellen olfactorische informatie kunnen herstellen "
     "uit natuurlijke taal. Zij evalueerden bijna 200 configuraties van drie generaties "
     "taalmodellen en concludeerden dat encoder-modellen (BERT-gebaseerd) het meest geschikt "
     "zijn voor olfactorische classificatietaken, terwijl grotere decoder-modellen (GPT-4o) "
     "beter presteren bij zero-shot olfactorische retrieval. Deze bevinding ondersteunt onze "
     "keuze voor een encoder-gebaseerd embeddingmodel als retrieval-component.")
h2("D. Positionering van Dit Onderzoek")
para("Tabel I vat de positionering van dit project samen ten opzichte van bestaand werk.")
para("TABEL I. Positionering ten opzichte van bestaand werk")
table(["Werk", "Methode", "Geurdomein", "Natural language input", "Interactief prototype"],
      [["Demarqui [7]", "Cosine similarity (BoW)", "ja", "nee", "nee"],
       ["Gonzales Martinez [8]", "Bayesiaans model", "ja", "nee", "nee"],
       ["Lewis et al. [4]", "RAG (generiek)", "nee", "ja", "nee"],
       ["Yousefi Maragheh et al. [9]", "Agentic RAG", "nee", "ja", "nee"],
       ["Dit project", "RAG + SBERT/BGE", "ja", "ja", "ja"]])
para("Dit project combineert als eerste een RAG-pipeline met sentence embeddings specifiek voor "
     "het geurdomein, met een werkend chatbot-prototype dat natural language input verwerkt.")

# ============================================================ III. DATA
h1("III. Data")
h2("A. Bronnen en omvang")
para("De dataset is samengesteld uit drie webbronnen, die elk via eigen scrapers zijn verzameld "
     "(zie Sectie III-C):")
bullet("Fragrantica master (perfumes_table.csv) - 84.144 parfums met zeven velden: rating, "
       "notes, designer, reviews, description, url, title. Dit is de hoofdbron en bevat de "
       "rijkste tekstuele signalen (geurnoten en gebruikersreviews).")
bullet("Parfumo-verrijking - een aanvullende scrape van Parfumo.com om ontbrekende ratings (en "
       "waar nodig beschrijvingen/reviews) in te vullen voor parfums die in Fragrantica geen "
       "rating hadden.")
bullet("Luckyscent (luckyscent_fragrances.csv) - circa 3.000 parfums met een afwijkend schema "
       "dat schone gender- en style-tags bevat.")
h2("B. Datakwaliteit en verkenning (EDA)")
para("Een exploratieve data-analyse (fragrantica_data_eda.ipynb) bracht enkele bepalende "
     "eigenschappen aan het licht die de rest van het project hebben gestuurd:")
bullet("Extreme rating-sparsity. Van de 84.144 parfums in de master had slechts 2.474 (ca. 3%) "
       "een rating. Voor de overige 81.670 ontbrak deze. Dit is de datamatige bevestiging van "
       "het in Sectie I beschreven rating-paradox: het overgrote deel van de catalogus is, in "
       "een op ratings gebaseerd systeem, onzichtbaar. Dit motiveerde zowel de "
       "Parfumo-verrijking als de keuze voor een retrieval-aanpak die niet primair op rating "
       "sorteert.")
bullet("Geclusterde ratings. De aanwezige ratings hadden een gemiddelde van 3,99/5 met een lage "
       "standaardafwijking (0,51). Ratings zijn dus weinig discriminerend - een extra argument "
       "om relevantie semantisch te bepalen in plaats van via rating.")
bullet("Rijke tekstvelden. 54.033 parfums (ca. 64%) bevatten reviews, met gemiddeld 15,9 "
       "reviews per parfum. Descriptions waren vrijwel volledig aanwezig (slechts 8 ontbrekend). "
       "Reviews en beschrijvingen bevatten precies de impliciete vibe-taal (\"evening\", "
       "\"seductive\", \"fresh\", \"lasts all day\") die queries als \"spicy for a date\" moeten "
       "matchen.")
bullet("Afleidbaar geslacht. gender zit niet als veld in Fragrantica maar in de title (\"... "
       "for women\"). Hieruit afgeleid: 36.616 unisex, 33.131 women, 14.389 men, 8 unknown. Dit "
       "geslacht is later als metadata-filter gebruikt.")
bullet("Luckyscent-kwaliteitsissues. De Luckyscent-EDA (luckyscent_eda.ipynb) toonde aan dat "
       "het brand-veld vrijwel overal de placeholder \"Brands\" bevatte (parsing-fout in de "
       "scraper) en dat price grotendeels \"$\" of \"$?\" was, dus onbruikbaar als numerieke "
       "feature. Luckyscent heeft bovendien geen reviews. De bruikbare signalen waren notes en "
       "style.")
para("Figuur 1 (fragrantica_fill_progress.png) visualiseert de status van de Fragrantica-dataset "
     "en de voortgang van de Parfumo-verrijking - aanvullende EDA-figuren (missing-values per "
     "kolom, ratingverdeling, top-20 geurnoten, correlatiematrix) zijn opgenomen in de "
     "bijbehorende notebooks.")
para("Figuur 1. Status van de Fragrantica-dataset en voortgang van de Parfumo-verrijking. "
     "[FIGUUR HIER INVOEGEN]")
h2("C. Dataverzameling en verrijking")
para("Het verzamelen van de data was technisch het meest arbeidsintensieve onderdeel, conform de "
     "observatie dat data cleaning vaak het meest tijdsintensieve deel van een "
     "datawetenschapsproject is.")
para("Parfumo-verrijking (webscrape_missing_parfumo_data.py). Omdat ca. 97% van de ratings "
     "ontbrak, hebben we deze niet weggegooid maar ingevuld via Parfumo. Dit bracht meerdere "
     "uitdagingen met zich mee, die expliciet zijn aangepakt:")
bullet("URL-reconstructie zonder API. Parfumo biedt geen API. Geldige profiel-URLs zijn "
       "gereconstrueerd uit designer en title (gender-suffixen verwijderen, merknaam uit de "
       "titel strippen, twee slug-varianten genereren omdat Parfumo inconsistent "
       "hyphenated-lowercase en Capitalized_Underscore gebruikt).")
bullet("Anti-bot-omzeiling. Fragrantica en Parfumo gebruiken Cloudflare-achtige bescherming. We "
       "gebruikten curl_cffi om de TLS-fingerprint van echte browsers na te bootsen, met rotatie "
       "van browserprofielen (Chrome/Edge/Safari) elke 15-35 requests, bijpassende HTTP-headers, "
       "en een session warmup (eerst de homepage bezoeken voor cookies).")
bullet("Mensachtige timing. Vertragingen volgen een random verdeling met lange (\"afgeleide "
       "gebruiker\") en korte pauzes, in plaats van een vast interval.")
bullet("Robuustheid. HTTP 429 triggert een oplopende backoff. Elk verzoek heeft een timeout en "
       "wordt tot 3x geprobeerd. Parfumo's \"valse 200\" (paginatitel \"404\") wordt "
       "gedetecteerd, voortgang wordt elke 50 rijen gecheckpoint, zodat een onderbroken scrape "
       "hervatbaar is.")
bullet("Werkverdeling. Met split_dataset.py is de set ontbrekende rijen in drie gelijke delen "
       "(missing_part1/2/3) opgesplitst, zodat drie teamleden parallel konden scrapen.")
para("Parfumo-ratings staan op een 0-10-schaal en zijn naar de Fragrantica-schaal (0-5) "
     "genormaliseerd via normalized_rating x 5, zodat alle ratings in dezelfde ruimte vallen.")
para("Luckyscent-scraper (luckyscent_scraper.ipynb). Een aparte BeautifulSoup-scraper pagineert "
     "door de productcatalogus en haalt naam, merk, notes, style, beschrijving, prijs, "
     "concentratie, gender en land op, met checkpoints elke 100 producten.")
h2("D. Annotaties")
para("Er zijn geen handmatige labels aangemaakt. Het systeem leunt op zwakke, afgeleide "
     "annotaties: het uit de titel afgeleide geslacht, de gestructureerde notes-lijsten, en de "
     "impliciete sentiment-/contextsignalen in description en reviews. Voor de evaluatie is een "
     "lijst van expected notes per testquery handmatig opgesteld (Sectie IV-H) en is een "
     "held-out testset (20%) afgesplitst.")

# ============================================================ IV. METHODE
h1("IV. Methode / Aanpak")
h2("A. Systeemarchitectuur")
para("De RAG-pipeline bestaat uit vier stappen:")
mono("Gebruiker typt beschrijving (natural language)\n"
     "        ->\n"
     "  Embeddingmodel  -> query-vector\n"
     "        ->\n"
     "  Vectordatabase  -> top-K meest vergelijkbare parfums (cosine)\n"
     "        ->\n"
     "  Filtering / re-ranking (gender, min_rating, rating-weight)\n"
     "        ->\n"
     "  Antwoordlaag (template of LLM) -> aanbeveling in natuurlijke taal")
para("Kort over de kernbegrippen: een embeddingmodel zet een stuk tekst om in een embedding - "
     "een vector, oftewel een lange rij getallen - zo gekozen dat teksten met een vergelijkbare "
     "betekenis dicht bij elkaar liggen. Zoeken komt dan neer op het vinden van de dichtstbijzijnde "
     "vectoren bij de query-vector, gemeten met cosine-gelijkenis (een getal tussen 0 en 1: hoe "
     "hoger, hoe meer de betekenis overeenkomt). De vectordatabase (hier ChromaDB) bewaart die "
     "vectoren en doet die nearest-neighbour-zoekopdracht razendsnel.")
h2("B. Data-integratie pipeline")
para("Voordat geindexeerd kon worden, zijn de drie bronnen samengevoegd tot een dataset "
     "(fragrantica_vectordb2.ipynb):")
bullet("Mergen op url - de Parfumo-verrijkte batches worden gededupliceerd (per url de rij met "
       "de minste lege cellen) en gebruikt om NaN-cellen in de master te overschrijven.")
bullet("Dedupliceren op (title, designer): dit vangt Fragrantica's dubbele listings (hetzelfde "
       "parfum onder twee URL-IDs). Behouden wordt de rij met de meeste gevulde cellen (meestal "
       "die met reviews).")
bullet("Normaliseren van Luckyscent naar het gedeelde schema en concatenatie.")
bullet("Kwaliteitsfilter - alleen rijen met zowel reviews als notes worden behouden, omdat "
       "reviews het sterkste vibe-signaal leveren en notes onmisbaar zijn voor de embedding. "
       "Luckyscent valt door dit filter (geen reviews) en wordt in de gefilterde index niet "
       "meegenomen.")
para("Embedding-tekst. Per parfum wordt een tekst-blob opgebouwd die expliciete features en "
     "impliciete vibes combineert: title + \"Brand:\" + designer + \"Gender:\" + gender + "
     "(\"Style:\" + style) + \"Notes:\" + notes + description + \"User reviews:\" + reviews. "
     "Reviews worden eerst opgeschoond: HTML-entities, URLs en \"show more\"-artefacten worden "
     "verwijderd, herhaalde leestekens ingekort, en - belangrijk - reviews die in 3 of meer "
     "verschillende parfums voorkomen worden gemarkeerd als scraper bleed-through en "
     "weggefilterd. Per parfum worden maximaal 5 reviews (langste eerst) tot 1.500 tekens "
     "meegenomen, zodat het meest informatieve signaal binnen het tekstbudget blijft.")

h2("C. Wat we hebben geprobeerd: van prototype naar productie")
para("Het systeem is iteratief tot stand gekomen: er zijn drie implementaties daadwerkelijk "
     "gebouwd en getest. De eerste twee zijn volwaardige, geevalueerde pogingen geweest - geen "
     "weggegooide schetsen, maar werkende notebook-prototypes die de haalbaarheid aantoonden en "
     "de bouwstenen (datapijplijn, embedding-tekst, query-expansie, gender-filtering) leverden. "
     "Op basis van hun beperkingen is vervolgens het uiteindelijke productiesysteem ontwikkeld, "
     "dat als enige in productie draait omdat het aanmerkelijk beter en robuuster presteert.")
para("Implementatie 1 - ChromaDB + all-MiniLM-L6-v2 (fragrantica_chatbot.ipynb). De volledige "
     "master (84.144 parfums) is geembed met all-MiniLM-L6-v2 (384-dim) en opgeslagen in een "
     "persistente ChromaDB-collectie met cosine-afstand, in batches van 500. ChromaDB "
     "ondersteunt een native WHERE-clause, waardoor het gender-filter direct in de query kan. "
     "Deze variant bevat ook de eerste chatbot- en evaluatielaag.")
para("Implementatie 2 - vectordb2 (Kagi) + fastembed BGE-small (fragrantica_vectordb2.ipynb). "
     "Een lichtgewicht alternatief: vectordb2 is een minimale vector store die standaard "
     "fastembed met het BGE-small-model gebruikt - sneller dan sentence-transformers en hoger "
     "scorend op MTEB dan MiniLM [11], [12]. Omdat vectordb2 geen WHERE-clause heeft, worden "
     "meer kandidaten opgehaald (top_k x 10) en gebeurt het filteren (gender, min_rating) en "
     "re-ranken achteraf in pandas.")
para("Beide verkennende stacks deelden dezelfde beperkingen: ze gebruikten een generiek "
     "embeddingmodel dat het geurdomein niet kent, kenden geen personalisatie, konden negatie "
     "(\"geen zoet\") niet hard afdwingen, en bleven notebook-prototypes zonder echte "
     "applicatie. Die beperkingen vormden de directe aanleiding voor een derde, geintegreerde "
     "implementatie - het uiteindelijke productiesysteem dat de verkennende notebooks samenbrengt "
     "en op elk van deze punten verbetert:")
para("Implementatie 3 - Fragrance Advisor: fine-tuned BGE-base + hybride recommender + webapp "
     "(recommender.py, app.py, build_collections.py). Dit systeem vervangt het generieke "
     "embeddingmodel door een zelf op het geurdomein fine-getuned model (Sectie IV-E), splitst "
     "de index in een notes- en een reviews-collectie (Sectie IV-F), en combineert drie "
     "retrieval-signalen - noten-vector, review-vector en een lexicale BM25-pass - via "
     "coverage-aware Reciprocal Rank Fusion, met personalisatie (smaakprofiel) en "
     "deterministische harde constraints (negatie/uitsluiting). Daarbovenop staat een "
     "conversationele webapplicatie met een verwisselbare LLM-laag (Sectie IV-G). De "
     "retrieval-core (recommender.py) wordt door zowel de evaluatie-notebook als de webapp "
     "gebruikt, zodat lab en productie exact dezelfde ranking draaien.")
para("Concreet verbetert Implementatie 3 op vier punten ten opzichte van de notebook-pogingen: "
     "(1) een domein-specifiek getraind embeddingmodel in plaats van een generiek model "
     "(Sectie IV-E), (2) hybride retrieval die ook review- en lexicaal signaal benut in plaats "
     "van enkel vector-similarity, (3) personalisatie via een smaakprofiel, en (4) "
     "deterministische afhandeling van negatie en uitsluiting. Samen maken deze het systeem "
     "zowel relevanter als betrouwbaarder; daarom is dit eindsysteem uitgerold. De twee "
     "verkennende metrieken (P@5 en de held-out Recall@10) meten verschillende dingen en zijn "
     "niet een-op-een vergelijkbaar - de winst van Implementatie 3 zit vooral in domein-fit, "
     "robuustheid en bruikbaarheid, niet in een enkel hoger getal.")
para("Tabel II vat de drie implementaties samen.")
para("TABEL II. De drie implementaties")
table(["Aspect", "Impl. 1 (ChromaDB)", "Impl. 2 (vectordb2)", "Impl. 3 (Fragrance Advisor)"],
      [["Embeddingmodel", "all-MiniLM-L6-v2 (384-dim)", "BGE-small (fastembed)",
        "fine-tuned BGE-base (768-dim)"],
       ["Retrieval", "enkel vector", "enkel vector",
        "notes-vector + reviews-vector + BM25 -> RRF"],
       ["Personalisatie", "nee", "nee", "ja (Rocchio-smaakprofiel)"],
       ["Negatie ('geen zoet')", "nee", "nee", "ja (deterministische harde filter)"],
       ["Filter op gender", "native WHERE", "post-hoc", "native WHERE"],
       ["Antwoordlaag", "template / OpenAI", "template / LLM-hook",
        "conversationele LLM (Ollama/OpenAI), grounded"],
       ["Interface", "ipywidgets", "notebook", "FastAPI-webapp"]])

h2("D. Belangrijke keuzes en parameters")
bullet("Re-ranking. De eindscore is (1 - rating_weight)*similarity + rating_weight*(rating/5). "
       "In Implementatie 2 is rating_weight = 0,2; in het productiesysteem 0,15. Zo wint bij "
       "gelijke relevantie het hoger gerate parfum, zonder dat rating de relevantie domineert - "
       "een bewuste demping van het rating-paradox.")
bullet("Query-expansie. Korte queries hebben weinig recall. Een handmatig samengestelde "
       "VIBE_HINTS-mapping breidt sleutelwoorden uit met gangbare noten/vibes (bijv. date -> "
       "\"evening seductive sensual warm intimate\", spicy -> \"cinnamon pepper saffron cardamom "
       "clove ginger\") voor het embedden.")
bullet("Filters. gender en min_rating zijn beschikbaar als harde filters; uitgesloten noten "
       "(Sectie IV-F) eveneens.")
bullet("Blend-gewicht. alpha = 0,5 weegt de gespreks-/zoektekst even zwaar als het "
       "smaakprofiel.")

h2("E. Fine-tuning van het embeddingmodel")
para("Een generiek embeddingmodel kent het geurdomein niet: het weet niet dat een review die "
     "over \"een warme avond, kaneel en leer\" praat, hoort bij een parfum met die noten. "
     "Daarom is een eigen bi-encoder getraind (fine-tuning/finetune_embeddings.py).")
bullet("Opzet. Een bi-encoder - een model dat de zoektekst en het document elk apart in een "
       "vector omzet, zodat ze daarna razendsnel te vergelijken zijn - is verder getraind vanuit "
       "BAAI/bge-base-en-v1.5 (dat 768-dimensionale vectoren maakt). De trainingsvoorbeelden zijn "
       "(review -> notes-passage)-paren: de review is de 'query' (de zoekkant) en de bijbehorende "
       "parfumtekst (\"title. Brand: ... Notes: ...\") is de 'positive', oftewel het juiste "
       "antwoord dat dicht bij de review moet komen te liggen. Uit de data kwamen 169.756 zulke "
       "paren; na het toevoegen van tegenvoorbeelden (zie hieronder) blijven 140.749 "
       "trainings-triplets over.")
bullet("Hard negatives (lastige tegenvoorbeelden). Een model leert het scherpst van bijna-goede "
       "fouten. Naast willekeurige tegenvoorbeelden uit dezelfde trainingsbatch (in-batch "
       "negatives) is daarom per paar een 'near-miss' gezocht: een parfum dat sterk op het juiste "
       "lijkt maar het net niet is (efficient opgezocht met FAISS). Zo leert het model de fijne "
       "verschillen tussen gelijkende parfums - de grootste kwaliteitswinst voor retrieval.")
bullet("Trainingsdoel en geheugen. De gebruikte loss (CachedMultipleNegativesRankingLoss) trekt "
       "de juiste paren naar elkaar toe en duwt de tegenvoorbeelden weg. Een geheugentruc "
       "(GradCache) maakt het mogelijk om met veel tegenvoorbeelden tegelijk te trainen "
       "(effectieve batch van 128) op een laptop-GPU met beperkt geheugen, door in kleine brokjes "
       "van 16 te rekenen.")
bullet("Lekkagevrije evaluatie. Een by-perfume split zet 1.000 hele parfums apart: zowel hun "
       "reviews als hun parfumtekst zijn tijdens het trainen nooit gezien. Zonder die scheiding "
       "zou het model het antwoord kunnen 'afkijken' (data-lekkage) en zouden de cijfers "
       "geflatteerd zijn. Elke epoch meet een Information-Retrieval-evaluator de Recall@10 (zie "
       "Sectie V-A voor de uitleg van die maat); het beste checkpoint wordt bewaard als "
       "models/bge-fragrance-v2.")
bullet("Resultaat. Het fine-tunen verhoogde de held-out Recall@10 van 0,42 (ongetraind) naar "
       "0,63 (zie Sectie V-A).")

h2("F. Hybride retrieval en personalisatie")
para("De productie-recommender (recommender.py) combineert drie complementaire signalen en "
     "personaliseert per gebruiker.")
bullet("Split-collecties (build_collections.py). De index is opgesplitst in twee "
       "ChromaDB-collecties: fragrances_notes met een document per parfum (alle 85.122) en "
       "fragrances_reviews met een document per losse recensie (met de parfum-URL als label). De "
       "reden voor die splitsing: een parfum heeft een vaste notenlijst maar vaak tientallen "
       "reviews; door reviews apart te indexeren telt elk ervaringsverhaal mee zonder de korte "
       "notes-tekst te overspoelen. Reviews worden opgeschoond, recensies die door een "
       "scrape-fout aan 3 of meer parfums gekoppeld zaten (scraper bleed-through) verwijderd, en "
       "per parfum maximaal 6 reviews bewaard. Van de 85.122 parfums hebben er 53.820 (circa 63%) "
       "bruikbare reviews.")
bullet("Drie zoekrondes ('passes') gecombineerd. Voor elke query draaien drie zoekmethoden "
       "naast elkaar. (1) Een notes-vectorpass: semantisch zoeken (op betekenis) in de "
       "notenlijst van alle parfums. (2) Een reviews-vectorpass: semantisch zoeken in de losse "
       "reviews, waarna elke gevonden review wordt teruggekoppeld naar het parfum waar hij bij "
       "hoort. (3) Een BM25-pass: BM25 is een klassieke trefwoord-zoekmethode (geen AI) die "
       "parfums rangschikt op hoe vaak en hoe kenmerkend de zoekwoorden letterlijk in hun noten, "
       "titel en merk voorkomen - dit vangt het geval waarin de gebruiker een specifieke noot "
       "noemt ('oud'), waar zoeken op betekenis te vaag kan zijn. De drie ranglijsten worden "
       "samengevoegd met Reciprocal Rank Fusion (RRF): elk parfum krijgt punten op basis van zijn "
       "positie in elke lijst, zodat parfums die in meerdere lijsten hoog staan bovenaan "
       "eindigen. Omdat de reviews-lijst niet alle parfums dekt (circa 37% heeft geen reviews), "
       "geeft een fairness-multiplier de parfums zonder reviews een duwtje, zodat de "
       "niche-catalogus niet structureel onder de mainstream parfums zakt.")
bullet("Personalisatie (Rocchio-methode). De gebruiker beoordeelt een paar bekende parfums "
       "(1-5). Hieruit wordt een 'smaakvector' opgebouwd volgens de klassieke Rocchio-techniek "
       "uit information retrieval: de vectoren van geliefde parfums worden opgeteld en die van "
       "afgekeurde afgetrokken (gewicht = cijfer - 3, dus een 5 trekt sterk aan, een 1 stoot "
       "sterk af). Die smaakvector wordt gemengd met de zoekvector; de parameter alpha bepaalt de "
       "verhouding tussen 'wat de gebruiker nu vraagt' en 'wat hij doorgaans lekker vindt'.")
bullet("Deterministische negatie (harde uitsluiting). Een verbod hoort niet in de embedding: "
       "\"geen zoet\" embedt het systeem juist richting zoet, want het woord 'zoet' staat er nu "
       "eenmaal. Daarom worden negaties met vaste regels (niet via het AI-model) uit de tekst "
       "gehaald: parse_negations herkent constructies als \"no/avoid/hate X\", en EXCLUDE_SYNONYMS "
       "breidt een categoriewoord uit naar de bijbehorende noten (sweet -> vanille, karamel, "
       "honing, ...). Kandidaten met zo'n noot, merk of naam worden na het ophalen hard "
       "weggefilterd, zodat een uitgesloten geur structureel nooit in de aanbeveling kan "
       "belanden.")
bullet("Faithfulness-check (controle op verzinsels). hallucinated_notes controleert achteraf of "
       "de gegenereerde tekst geen noten noemt die in geen van de aanbevolen parfums voorkomen - "
       "een geautomatiseerd vangnet tegen 'hallucinatie' (een taalmodel dat plausibel klinkende "
       "maar verzonnen details toevoegt).")

h2("G. Webapplicatie en gespreksinterface")
para("Het prototype is uitgewerkt tot een werkende webapplicatie (app.py, FastAPI + "
     "static/index.html) met een tweetraps-gespreksflow en een verwisselbare LLM-backend.")
bullet("Tweetraps-LLM. (1) Een advisor voert eerst de intake: hij stelt maximaal drie "
       "verhelderende vragen (gelegenheid, seizoen, intensiteit, te vermijden noten, gender) en "
       "geeft een onzichtbaar signaalwoord [RECOMMEND] af zodra hij genoeg weet. De advisor heeft "
       "geen toegang tot de catalogus en mag dus nooit zelf parfums noemen; een stukje code (een "
       "'sanitizer') verwijdert eventuele door het model verzonnen lijstjes. (2) Pas daarna "
       "genereert een tweede, 'grounded' stap de aanbeveling. Grounded betekent dat het taalmodel "
       "uitsluitend mag praten over de parfums die de zoekmachine heeft aangeleverd en alleen "
       "noten mag noemen die in hun notenlijst staan - het mag niets bijverzinnen.")
bullet("Verwisselbare backend. Met een regel schakelt het systeem tussen Ollama (lokaal "
       "qwen2.5:7b, gratis) en OpenAI. Negaties uit het gesprek worden als harde filter "
       "doorgegeven, niet via de prompt.")
bullet("Endpoints. /api/search (autocomplete over de catalogus voor het opbouwen van het "
       "smaakprofiel), /api/chat (advisor + aanbeveling), en de frontend op /.")

h2("H. Evaluatie")
para("Evaluatie is een kernonderdeel van dit project. Er bestaat geen publieke ground-truth voor "
     "\"de juiste geur bij een beschrijving\", en geuraanbeveling is intrinsiek subjectief. "
     "Daarom is niet op een enkele score vertrouwd, maar zijn vijf complementaire, deels "
     "geautomatiseerde methoden gebruikt - elk met een eigen doel, eigen protocol en eigen "
     "beperking (Tabel III). Twee zijn kwantitatief (Recall@10, Precision@5), een is een "
     "geautomatiseerde betrouwbaarheidscontrole (faithfulness), en twee zijn analytisch/"
     "kwalitatief.")
para("TABEL III. Overzicht van de evaluatiemethoden")
table(["Methode", "Wat het meet", "Toegepast op"],
      [["Held-out Recall@1/5/10 (IR-evaluator)", "Retrieval-generalisatie van het embeddingmodel",
        "fine-tuned BGE-base"],
       ["Precision@5 (note-overlap)", "Thematische relevantie van de top-5", "Impl. 1"],
       ["Gem. cosine similarity", "Retrieval-zekerheid per query", "Impl. 1"],
       ["Faithfulness-check (hallucinated_notes)", "Trouw van de LLM-tekst aan de kandidaten",
        "Impl. 3"],
       ["Kwalitatieve inspectie", "Negatie, personalisatie, rating-ontkoppeling", "Impl. 3"]])
para("1) Retrieval-generalisatie (held-out Recall@10). Dit is de strengste en meest "
     "kwantitatieve toets, gebruikt om het fine-tunen te sturen (Sectie IV-E). De opzet is "
     "lekkagevrij: via een by-perfume split worden 1.000 parfums volledig buiten de training "
     "gehouden - zowel hun reviews als hun notes-passage zijn tijdens het trainen nooit gezien. "
     "De zoekruimte (corpus) bestaat uit die 1.000 relevante passages plus 5.000 willekeurige "
     "train-passages als afleiders (6.000 in totaal); de query is de langste held-out review van "
     "elk parfum. Een Information-Retrieval-evaluator (sentence-transformers) meet per epoch "
     "Recall@1/5/10, MRR@10, NDCG@10 en MAP@10, en het beste checkpoint wordt teruggeladen. Kort "
     "uitgelegd: Recall@10 is het percentage gevallen waarin het juiste parfum in de top-10 "
     "staat; MRR (Mean Reciprocal Rank) en NDCG belonen daarbovenop een hoge positie van het "
     "juiste antwoord - ze tellen een treffer op plek 1 zwaarder dan een op plek 9.")
para("2) Baseline-vergelijking. Exact dezelfde evaluator draait op het ongetrainde BGE-base "
     "model vóór en op de fine-tuned versie ná de training, zodat de verbetering op identieke "
     "data en een identiek protocol wordt afgelezen. De fine-tuned versie is alleen behouden "
     "omdat zij op deze held-out taak beter scoorde (Sectie V-A).")
para("3) Thematische relevantie (Precision@5). Voor vijf thematische queries is vooraf een lijst "
     "expected notes gedefinieerd; een resultaat telt als relevant wanneer ten minste een "
     "verwachte noot in de opgehaalde geurnoten voorkomt, en P@5 is het aandeel relevante "
     "resultaten in de top-5. Per query is daarnaast de gemiddelde cosine similarity berekend als "
     "indicatie van retrieval-zekerheid.")
para("4) Faithfulness-/hallucinatiecontrole. Omdat de LLM-laag tekst genereert, is hallucinatie "
     "geautomatiseerd meetbaar gemaakt: hallucinated_notes markeert noot-woorden in de "
     "LLM-output die wel in het globale noten-vocabulaire bestaan maar in geen van de aanbevolen "
     "parfums voorkomen. Zo is trouw aan de opgehaalde kandidaten een controleerbare eigenschap "
     "in plaats van een aanname.")
para("5) Kwalitatieve inspectie. Op vrije queries (Engels en Nederlands) is gecontroleerd of "
     "negatie correct hard wordt afgedwongen, of personalisatie de resultaten zichtbaar "
     "verschuift, en of de aanbevelingen loskomen van de community-rating.")
para("Data-splitdiscipline. Voor de notebook-retrieval is een 80/20 train/test-split (seed 42) "
     "gebruikt en voor het fine-tunen de strengere by-perfume held-out split. Bij de "
     "notebook-varianten dient de split niet tegen overfitting - de embeddings komen uit een "
     "frozen model - maar voorkomt het dat een testquery exact de referentierij ophaalt.")
para("Beperkingen van de evaluatie (zie ook Sectie VI-B). P@5 gebruikt note-overlap als proxy en "
     "een kleine queryset; Recall@10 hanteert het parfum van de review als enige relevante item, "
     "een streng maar smal relevantiebegrip; en een formele gebruikersstudie met echte "
     "tevredenheidsoordelen ontbreekt nog.")

# ============================================================ V. RESULTATEN
h1("V. Resultaten")
h2("A. Retrieval-generalisatie van het fine-tuned model")
para("Dit is het belangrijkste kwantitatieve resultaat. Op de volledig held-out set (1.000 "
     "ongeziene parfums + 5.000 afleider-passages, review-als-query) behaalt bge-fragrance-v2 de "
     "scores in Tabel IV.")
para("TABEL IV. Held-out retrieval van het fine-tuned model (review-als-query, 6.000 kandidaten)")
table(["Metriek", "Score", "Interpretatie"],
      [["Recall@1 / Accuracy@1", "0,464", "juiste parfum direct op plek 1"],
       ["Recall@5", "0,586", "juiste parfum in de top-5"],
       ["Recall@10", "0,633", "juiste parfum in de top-10"],
       ["MRR@10", "0,515", "gemiddelde reciproke rang"],
       ["NDCG@10", "0,543", "rang-gewogen kwaliteit"]])
para("In bijna twee op de drie gevallen (Recall@10 = 0,63) staat het exact juiste parfum in de "
     "top-10 wanneer alleen een ongeziene review als zoekopdracht wordt gegeven - een streng "
     "criterium, want er is precies een correct doel tussen 6.000 kandidaten. In bijna de helft "
     "van de gevallen staat het zelfs meteen op plek 1 (Recall@1 = 0,46). De relatief hoge "
     "MRR@10 (0,52) en NDCG@10 (0,54) laten zien dat wanneer het juiste parfum wordt gevonden, "
     "het doorgaans hoog in de lijst staat in plaats van ergens onderaan.")
para("Effect van het fine-tunen (Tabel V). Om te bewijzen dat het trainen echt iets toevoegt, is "
     "dezelfde evaluator vóór de training op het ongetrainde BGE-base model gedraaid en erna op "
     "de fine-tuned versie, op identieke held-out data en protocol. De Recall@10 steeg daarbij "
     "van 0,42 naar 0,63 - een toename van ruim 50% relatief. Het juiste parfum komt na het "
     "fine-tunen dus veel vaker in de top-10. Die winst komt vooral van de hard-negative mining, "
     "die het model de fijne onderscheidingen tussen gelijkende parfums leert. Het beste "
     "resultaat lag al na de eerste epoch (de tweede gaf met Recall@10 = 0,630 geen verbetering "
     "t.o.v. 0,633), dus is het eerste checkpoint bewaard. De training duurde circa 1,4 uur op "
     "een laptop-GPU (NVIDIA RTX 5050).")
para("TABEL V. Effect van het fine-tunen op de held-out Recall@10")
table(["Model", "Recall@10"],
      [["BGE-base (ongetraind, baseline)", "0,416"],
       ["bge-fragrance-v2 (fine-tuned)", "0,633"]])
h2("B. Thematische relevantie (Implementatie 1)")
para("De note-overlap-evaluatie op vijf thematische queries gaf een gemiddelde Precision@5 van "
     "0,92: gemiddeld zijn 4,6 van de 5 top-resultaten thematisch relevant (Tabel VI). De "
     "gemiddelde similarity (0,55-0,70) was hoger bij goed-gelexicaliseerde thema's (floral, "
     "0,70) dan bij abstractere (woody/oud, 0,56) - consistent met de schaarste van "
     "olfactorische taal [10]. Deze maat dekt thematische passendheid, maar niet of exact het "
     "beste parfum bovenaan staat; daarom is hij complementair aan de Recall@10 uit Sectie V-A.")
para("TABEL VI. Precision@5 en gemiddelde similarity per query (Impl. 1)")
table(["Query", "Relevant (van 5)", "Gem. similarity", "P@5"],
      [["fresh citrus summer beach", "5", "0,582", "1,00"],
       ["dark woody oud oriental", "4", "0,555", "0,80"],
       ["sweet vanilla warm gourmand", "4", "0,551", "0,80"],
       ["floral rose jasmine feminine", "5", "0,701", "1,00"],
       ["spicy pepper saffron", "5", "0,599", "1,00"],
       ["Gemiddeld", "4,6", "0,598", "0,92"]])
h2("C. Systeemgedrag: faithfulness, negatie en personalisatie (Implementatie 3)")
para("Naast de retrieval-metrieken is het gedrag van het eindsysteem geevalueerd met de "
     "faithfulness-check en kwalitatieve inspectie (Sectie IV-H, methoden 4 en 5). De Fragrance "
     "Advisor draait end-to-end: de gebruiker beoordeelt enkele bekende parfums, voert een kort "
     "gesprek, en krijgt een gepersonaliseerde, in natuurlijke taal onderbouwde aanbeveling. De "
     "belangrijkste bevindingen:")
bullet("Negatie werkt. Door deterministisch filteren leidt \"iets fris, maar geen zoet\" niet "
       "langer tot zoete aanbevelingen - een bekend faalpunt van puur embedding-gebaseerde "
       "systemen, waar het woord \"zoet\" de query juist naar zoet trekt.")
bullet("Personalisatie verschuift de resultaten zichtbaar richting de smaak die uit de "
       "beoordeelde parfums spreekt, terwijl dislikes worden weggeduwd.")
bullet("Grounded antwoorden. De aanbeveltekst noemt uitsluitend parfums en noten uit de "
       "opgehaalde kandidaten; de faithfulness-check en de strikte prompt beperken "
       "hallucinatie.")
bullet("Rating-paradox bevestigd. Net als in de notebooks haalt het systeem regelmatig laag- of "
       "niet-gewaardeerde parfums naar boven wanneer die het best matchen - precies de geuren "
       "die in een op rating gesorteerd platform onzichtbaar blijven.")

# ============================================================ VI. DISCUSSIE
h1("VI. Discussie en Conclusie")
h2("A. Interpretatie")
para("De resultaten beantwoorden de hoofdonderzoeksvraag bevestigend: een RAG-gebaseerde "
     "pipeline kan gepersonaliseerde geuraanbevelingen uit vrije-tekstinput genereren. Het "
     "domein-specifiek fine-tunen tilt retrieval naar Recall@10 = 0,63 op een streng held-out "
     "criterium, en de hybride combinatie van noten-, review- en lexicaal signaal met "
     "deterministische constraints maakt het systeem robuuster dan de losse notebook-varianten. "
     "Belangrijker dan de cijfers is dat de aanbevelingen loskomen van populariteit: het systeem "
     "adresseert het rating-paradox dat traditionele filtering kenmerkt.")
h2("B. Wat ging goed, wat kon beter")
para("Sterk. De volledige keten staat: data uit drie bronnen samenvoegen, verrijken via Parfumo "
     "onder zware anti-bot-condities, een eigen embeddingmodel trainen, splitsen in "
     "notes/reviews, hybride fuseren met personalisatie en harde negatie, en dat alles ontsluiten "
     "via een conversationele webapp met een grounded LLM. Een gedeelde retrieval-core "
     "garandeert dat lab en productie identiek zijn.")
para("Beperkingen.")
bullet("Evaluatie deels via proxy. P@5 meet note-overlap en Recall@10 meet of een specifieke "
       "review zijn eigen parfum terugvindt - beide zijn nuttige maar indirecte maten voor echte "
       "gebruikerstevredenheid. Een formele gebruikersstudie (A/B) ontbreekt nog.")
bullet("LLM-afhankelijkheid. De gesprekskwaliteit hangt af van de gekozen backend; het lokale "
       "qwen2.5:7b wijkt soms van de instructie af, wat we deterministisch moeten afvangen.")
bullet("Datakwaliteit en dekking. Luckyscent heeft een kapot merk-veld en onbruikbare prijzen; "
       "de Parfumo-verrijking en de review-dekking (ca. 38% na strikte filtering) zijn niet "
       "volledig.")
bullet("Olfactorische taalschaarste [10] legt een fundamenteel plafond op puur taalgebaseerde "
       "matching, vooral bij abstracte queries.")
bullet("Meertaligheid (Nederlandse queries) werkt deels maar is niet systematisch geevalueerd; "
       "het fine-getunede model is Engelstalig.")
h2("C. Ethische overwegingen")
para("(1) Webscraping: de scrapers respecteren rate-limits en bootsen normaal gedrag na, maar "
     "het scrapen van Fragrantica/Parfumo raakt aan hun voorwaarden en aan het auteursrecht op "
     "reviews; voor productie is een licentie of officiele databron nodig. (2) Bias: "
     "community-reviews en -ratings bevatten populariteits- en taalbias (overwegend Engels), die "
     "in de embeddings doorwerken. (3) Sturing en transparantie: als commercieel platform kan het "
     "systeem koopgedrag sturen; uitlegbare, grounded aanbevelingen en het niet-verbergen van "
     "niche-opties zijn bewuste keuzes. (4) Privacy: gescrapete reviews zijn van echte gebruikers "
     "en worden uitsluitend geaggregeerd, niet herleidbaar, gebruikt.")
h2("D. Vervolgstappen")
bullet("Gebruikersstudie: een A/B-opzet met echte gebruikers en relevantielabels, zodat naast "
       "Recall@10 en P@5 ook tevredenheid wordt gemeten.")
bullet("Tuning van de fusie: review_weight, review_fairness, alpha en rating_weight systematisch "
       "optimaliseren op een gelabelde validatieset.")
bullet("Datacompleetheid: de Parfumo-verrijking afronden, de review-dekking vergroten en de "
       "Luckyscent-scraper repareren (echte merknaam, numerieke prijs).")
bullet("Meertalig model: een meertalige of Nederlandse fine-tune voor robuuste NL-input.")
bullet("Conversationeel geheugen: het smaakprofiel over sessies bewaren en de aanbeveling "
       "iteratief verfijnen.")
h2("E. Conclusie")
para("Dit onderzoek demonstreert een volledige RAG-pipeline voor het geurdomein, van "
     "dataverzameling tot een werkende, conversationele webapplicatie. Door een eigen "
     "embeddingmodel te fine-tunen (Recall@10 = 0,63 held-out), drie retrieval-signalen hybride "
     "te fuseren, te personaliseren met een smaakprofiel en harde constraints deterministisch te "
     "handhaven, levert het systeem relevante, uitlegbare en van populariteit losgekoppelde "
     "aanbevelingen. Daarmee toont het de meerwaarde van een vectorgebaseerde, gepersonaliseerde "
     "aanpak boven traditionele, op rating gebaseerde filtering, en legt het een onderbouwde "
     "basis voor verdere ontwikkeling door Fragheadsunited.")

# ============================================================ REFERENTIES
h1("Referenties")
refs = [
    "[1] Statista, \"Fragrances - Netherlands: Revenue and market data,\" 2024. [Online]. "
    "Available: https://www.statista.com/outlook/cmo/beauty-personal-care/fragrances/netherlands",
    "[2] WGSN, \"Consumer behaviour forecast: Scent identity and personalization trends,\" 2024.",
    "[3] Euromonitor International, \"Unlocking fragrance consumer trends in Western Europe,\" "
    "2024. [Online]. Available: "
    "https://www.euromonitor.com/article/unlocking-fragrance-consumer-trends-in-western-europe",
    "[4] P. Lewis et al., \"Retrieval-augmented generation for knowledge-intensive NLP tasks,\" "
    "in Proc. NeurIPS, vol. 33, 2020, pp. 9459-9474.",
    "[5] N. Reimers and I. Gurevych, \"Sentence-BERT: Sentence embeddings using Siamese "
    "BERT-networks,\" in Proc. EMNLP-IJCNLP, 2019, pp. 3982-3992.",
    "[6] F. Ricci, L. Rokach, B. Shapira, and P. B. Kantor, Recommender Systems Handbook. New "
    "York, NY: Springer, 2011.",
    "[7] R. Demarqui, \"Perfume recommendation system based on Fragrantica data,\" GitHub "
    "repository, 2022. [Online]. Available: https://github.com/rdemarqui/perfume_recommendation",
    "[8] L. Gonzales Martinez, \"Fragrance recommendation engine using Bayesian hierarchical "
    "models and Jungian archetypes,\" M.S. thesis, Dept. Data Sci., Univ. San Francisco, 2024.",
    "[9] H. Yousefi Maragheh et al., \"ARAG: An agentic RAG framework for personalized "
    "recommendations,\" arXiv preprint arXiv:2504.02620, 2025.",
    "[10] M. Kurfali, T. Horberg, and J. Ostling, \"Can language models smell? Investigating "
    "olfactory language understanding in language models,\" in Proc. ACL, 2024.",
    "[11] S. Xiao, Z. Liu, P. Zhang, and N. Muennighoff, \"C-Pack: Packed resources for general "
    "Chinese embeddings,\" arXiv preprint arXiv:2309.07597, 2023.",
    "[12] N. Muennighoff, N. Tazi, L. Magne, and N. Reimers, \"MTEB: Massive text embedding "
    "benchmark,\" in Proc. EACL, 2023, pp. 2014-2037.",
    "[13] Chroma, \"Chroma: the open-source AI application database,\" 2024. [Online]. "
    "Available: https://www.trychroma.com",
    "[14] S. Robertson and H. Zaragoza, \"The probabilistic relevance framework: BM25 and "
    "beyond,\" Found. Trends Inf. Retr., vol. 3, no. 4, pp. 333-389, 2009.",
    "[15] G. V. Cormack, C. L. A. Clarke, and S. Buttcher, \"Reciprocal rank fusion outperforms "
    "Condorcet and individual rank learning methods,\" in Proc. ACM SIGIR, 2009, pp. 758-759.",
    "[16] J. J. Rocchio, \"Relevance feedback in information retrieval,\" in The SMART Retrieval "
    "System, G. Salton, Ed. Englewood Cliffs, NJ: Prentice-Hall, 1971, pp. 313-323.",
    "[17] Qwen Team, \"Qwen2.5 technical report,\" arXiv preprint arXiv:2412.15115, 2024.",
]
for r in refs:
    para(r)

import time
out = "technisch_rapport.docx"
try:
    doc.save(out)
except PermissionError:
    out = f"technisch_rapport_{time.strftime('%H%M%S')}.docx"
    doc.save(out)
    print("(technisch_rapport.docx stond open in Word - opgeslagen onder nieuwe naam)")
print("Opgeslagen:", out)
